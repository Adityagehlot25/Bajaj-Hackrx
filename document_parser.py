"""
Document parser module for extracting text from PDF, DOCX, and .eml files
and splitting them into logical chunks.
"""

import email
import re
from email import policy
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    logger.warning("PyPDF2 not available. PDF parsing will be disabled.")

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX parsing will be disabled.")

class DocumentChunk:
    """Represents a chunk of text from a document."""
    
    def __init__(self, text: str, chunk_id: int, source_file: str, 
                 start_pos: int = 0, end_pos: int = 0, metadata: Dict[str, Any] = None):
        self.text = text.strip()
        self.chunk_id = chunk_id
        self.source_file = source_file
        self.start_pos = start_pos
        self.end_pos = end_pos
        self.metadata = metadata or {}
        self.word_count = len(text.split())

class DocumentParser:
    """Main class for parsing documents and extracting text chunks."""
    
    def __init__(self, min_chunk_words: int = 500, max_chunk_words: int = 1000):
        self.min_chunk_words = min_chunk_words
        self.max_chunk_words = max_chunk_words
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a file and return extracted text with chunks.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Dictionary containing extracted text, chunks, and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Determine file type and parse accordingly
        if file_path.suffix.lower() == '.pdf':
            return self._parse_pdf(file_path)
        elif file_path.suffix.lower() in ['.docx', '.doc']:
            return self._parse_docx(file_path)
        elif file_path.suffix.lower() == '.eml':
            return self._parse_eml(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")
    
    def _parse_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF file and extract text."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is required for PDF parsing. Install with: pip install PyPDF2")
        
        text_content = ""
        metadata = {"file_type": "pdf", "pages": 0}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
                        
        except Exception as e:
            raise Exception(f"Error parsing PDF: {e}")
        
        chunks = self._split_into_chunks(text_content, str(file_path))
        
        return {
            "file_path": str(file_path),
            "file_type": "pdf",
            "raw_text": text_content,
            "chunks": [self._chunk_to_dict(chunk) for chunk in chunks],
            "total_chunks": len(chunks),
            "total_words": len(text_content.split()),
            "metadata": metadata
        }
    
    def _parse_docx(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file and extract text."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing. Install with: pip install python-docx")
        
        text_content = ""
        metadata = {"file_type": "docx", "paragraphs": 0}
        
        try:
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    metadata["paragraphs"] += 1
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content += cell.text + "\n"
                            
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {e}")
        
        chunks = self._split_into_chunks(text_content, str(file_path))
        
        return {
            "file_path": str(file_path),
            "file_type": "docx",
            "raw_text": text_content,
            "chunks": [self._chunk_to_dict(chunk) for chunk in chunks],
            "total_chunks": len(chunks),
            "total_words": len(text_content.split()),
            "metadata": metadata
        }
    
    def _parse_eml(self, file_path: Path) -> Dict[str, Any]:
        """Parse .eml email file and extract text."""
        text_content = ""
        metadata = {"file_type": "eml"}
        
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                msg = email.message_from_file(file, policy=policy.default)
            
            # Extract email headers
            metadata.update({
                "subject": msg.get('Subject', 'No Subject'),
                "from": msg.get('From', 'Unknown Sender'),
                "to": msg.get('To', 'Unknown Recipient'),
                "date": msg.get('Date', 'Unknown Date')
            })
            
            # Add header information to text content
            text_content += f"Subject: {metadata['subject']}\n"
            text_content += f"From: {metadata['from']}\n"
            text_content += f"To: {metadata['to']}\n"
            text_content += f"Date: {metadata['date']}\n\n"
            
            # Extract email body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        body = part.get_payload(decode=True)
                        if body:
                            try:
                                text_content += body.decode('utf-8', errors='ignore')
                            except:
                                text_content += str(body)
            else:
                body = msg.get_payload(decode=True)
                if body:
                    try:
                        text_content += body.decode('utf-8', errors='ignore')
                    except:
                        text_content += str(body)
                        
        except Exception as e:
            raise Exception(f"Error parsing EML: {e}")
        
        chunks = self._split_into_chunks(text_content, str(file_path))
        
        return {
            "file_path": str(file_path),
            "file_type": "eml",
            "raw_text": text_content,
            "chunks": [self._chunk_to_dict(chunk) for chunk in chunks],
            "total_chunks": len(chunks),
            "total_words": len(text_content.split()),
            "metadata": metadata
        }
    
    def _split_into_chunks(self, text: str, source_file: str) -> List[DocumentChunk]:
        """
        Split text into logical chunks of 500-1000 words each.
        
        Args:
            text: The text to split
            source_file: Path to the source file
            
        Returns:
            List of DocumentChunk objects
        """
        if not text.strip():
            return []
        
        chunks = []
        
        # Split text into sentences
        sentences = self._split_into_sentences(text)
        
        current_chunk = ""
        current_word_count = 0
        chunk_id = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            
            # If adding this sentence would exceed max_chunk_words, finalize current chunk
            if current_word_count + sentence_words > self.max_chunk_words and current_chunk:
                if current_word_count >= self.min_chunk_words:
                    chunks.append(DocumentChunk(
                        text=current_chunk,
                        chunk_id=chunk_id,
                        source_file=source_file,
                        metadata={"word_count": current_word_count}
                    ))
                    chunk_id += 1
                    current_chunk = ""
                    current_word_count = 0
            
            current_chunk += sentence + " "
            current_word_count += sentence_words
        
        # Add the final chunk if it has content
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                text=current_chunk,
                chunk_id=chunk_id,
                source_file=source_file,
                metadata={"word_count": current_word_count}
            ))
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences using regex patterns."""
        # Clean up text first
        text = re.sub(r'\s+', ' ', text)  # Replace multiple whitespace with single space
        text = text.strip()
        
        # Split on sentence boundaries
        sentences = re.split(r'(?<=[.!?])\s+', text)
        
        # Filter out empty sentences and very short ones
        sentences = [s.strip() for s in sentences if s.strip() and len(s.split()) > 3]
        
        return sentences
    
    def _chunk_to_dict(self, chunk: DocumentChunk) -> Dict[str, Any]:
        """Convert DocumentChunk to dictionary."""
        return {
            "chunk_id": chunk.chunk_id,
            "text": chunk.text,
            "word_count": chunk.word_count,
            "source_file": chunk.source_file,
            "metadata": chunk.metadata
        }

def parse_document(file_path: str, min_chunk_words: int = 500, max_chunk_words: int = 1000) -> Dict[str, Any]:
    """
    Convenience function to parse a document and return chunks.
    
    Args:
        file_path: Path to the document file
        min_chunk_words: Minimum words per chunk (default: 500)
        max_chunk_words: Maximum words per chunk (default: 1000)
        
    Returns:
        Dictionary containing parsed document information and chunks
    """
    parser = DocumentParser(min_chunk_words, max_chunk_words)
    return parser.parse_file(file_path)
