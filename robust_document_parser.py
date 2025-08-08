"""
Robust Document Parser with Multiple PDF Libraries and Comprehensive Error Handling
Supports PyMuPDF (fitz), pdfplumber, PyPDF2 with fallback mechanisms
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import json
from datetime import datetime

# Configure detailed logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Try to import tiktoken for accurate token counting
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
    logger.info("tiktoken available for accurate token counting")
    # Use GPT-3.5-turbo encoding as a good general-purpose tokenizer
    TOKEN_ENCODER = tiktoken.get_encoding("cl100k_base")
except ImportError:
    TIKTOKEN_AVAILABLE = False
    logger.warning("tiktoken not available. Using approximate word-based token counting.")
    TOKEN_ENCODER = None

# PDF parsing libraries with fallback priority (best to worst)
PDF_LIBRARIES = []

try:
    import fitz  # PyMuPDF
    PDF_LIBRARIES.append(('pymupdf', fitz))
    logger.info("PyMuPDF (fitz) available for PDF parsing")
except ImportError:
    logger.warning("PyMuPDF not available. Install with: pip install PyMuPDF")

try:
    import pdfplumber
    PDF_LIBRARIES.append(('pdfplumber', pdfplumber))
    logger.info("pdfplumber available for PDF parsing")
except ImportError:
    logger.warning("pdfplumber not available. Install with: pip install pdfplumber")

try:
    import PyPDF2
    PDF_LIBRARIES.append(('pypdf2', PyPDF2))
    logger.info("PyPDF2 available for PDF parsing")
except ImportError:
    logger.warning("PyPDF2 not available. Install with: pip install PyPDF2")

if not PDF_LIBRARIES:
    logger.error("No PDF parsing libraries available! Install at least one: PyMuPDF, pdfplumber, or PyPDF2")

try:
    from docx import Document
    DOCX_AVAILABLE = True
    logger.info("python-docx available for DOCX parsing")
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Install with: pip install python-docx")


def estimate_tokens(text: str) -> int:
    """
    Estimate token count for text using tiktoken if available, otherwise use approximation.
    
    Args:
        text: Input text to count tokens for
        
    Returns:
        Estimated number of tokens
    """
    if not text:
        return 0
        
    if TIKTOKEN_AVAILABLE and TOKEN_ENCODER:
        try:
            return len(TOKEN_ENCODER.encode(text))
        except Exception as e:
            logger.warning(f"tiktoken encoding failed: {e}, falling back to approximation")
    
    # Fallback: approximate tokens as words * 1.3 (typical ratio for English text)
    word_count = len(text.split())
    return int(word_count * 1.3)


def split_text_recursively(text: str, max_tokens: int = 2000, min_tokens: int = 100) -> List[str]:
    """
    Recursively split text into chunks that don't exceed the maximum token count.
    
    This implements a multi-level splitting strategy:
    1. Split by double newlines (paragraphs/sections)
    2. If chunks are still too large, split by sentences
    3. If still too large, split by fixed character length at safe delimiters
    4. Ensure no chunk exceeds the token limit
    
    Args:
        text: Text to split
        max_tokens: Maximum tokens per chunk (default 2000 for safety margin)
        min_tokens: Minimum tokens per chunk to avoid tiny fragments
        
    Returns:
        List of text chunks, each under the token limit
    """
    if not text or not text.strip():
        return []
    
    text = text.strip()
    current_tokens = estimate_tokens(text)
    
    # If text is already small enough, return it as a single chunk
    if current_tokens <= max_tokens:
        if current_tokens >= min_tokens:
            return [text]
        else:
            # Text is too small to be a standalone chunk
            return [text]  # Return it anyway, will be combined later
    
    logger.debug(f"Splitting text with {current_tokens} tokens (max: {max_tokens})")
    
    # Level 1: Split by paragraphs (double newlines)
    paragraph_chunks = []
    paragraphs = re.split(r'\n\s*\n', text)
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
            
        para_tokens = estimate_tokens(para)
        
        if para_tokens <= max_tokens:
            paragraph_chunks.append(para)
        else:
            # Level 2: Split large paragraph by sentences
            sentence_chunks = _split_by_sentences(para, max_tokens, min_tokens)
            paragraph_chunks.extend(sentence_chunks)
    
    # Level 3: Combine small chunks and handle edge cases
    final_chunks = _combine_small_chunks(paragraph_chunks, max_tokens, min_tokens)
    
    # Validation: ensure no chunk exceeds token limit
    validated_chunks = []
    for chunk in final_chunks:
        chunk_tokens = estimate_tokens(chunk)
        if chunk_tokens > max_tokens:
            logger.warning(f"Chunk still too large ({chunk_tokens} tokens), applying character-based splitting")
            # Level 4: Character-based splitting as last resort
            char_chunks = _split_by_characters(chunk, max_tokens)
            validated_chunks.extend(char_chunks)
        else:
            validated_chunks.append(chunk)
    
    logger.info(f"Recursive splitting: {current_tokens} tokens → {len(validated_chunks)} chunks")
    return validated_chunks


def _split_by_sentences(text: str, max_tokens: int, min_tokens: int) -> List[str]:
    """Split text by sentences, respecting token limits."""
    # Enhanced sentence splitting regex that handles various punctuation
    sentence_endings = r'(?<=[.!?])\s+(?=[A-Z])'
    sentences = re.split(sentence_endings, text)
    
    chunks = []
    current_chunk = ""
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence:
            continue
        
        # Check if adding this sentence would exceed the token limit
        potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
        potential_tokens = estimate_tokens(potential_chunk)
        
        if potential_tokens <= max_tokens:
            current_chunk = potential_chunk
        else:
            # Save current chunk if it meets minimum requirements
            if current_chunk and estimate_tokens(current_chunk) >= min_tokens:
                chunks.append(current_chunk)
            
            # Check if single sentence is too large
            sentence_tokens = estimate_tokens(sentence)
            if sentence_tokens > max_tokens:
                logger.warning(f"Single sentence too large ({sentence_tokens} tokens), will be character-split")
                chunks.append(sentence)  # Will be handled by character splitting later
            else:
                current_chunk = sentence
    
    # Add remaining chunk
    if current_chunk:
        chunks.append(current_chunk)
    
    return chunks


def _split_by_characters(text: str, max_tokens: int) -> List[str]:
    """Split text by character count at safe delimiters as last resort."""
    chunks = []
    
    # Estimate characters per token (rough approximation)
    chars_per_token = len(text) / max(1, estimate_tokens(text))
    target_chars = int(max_tokens * chars_per_token * 0.9)  # 90% safety margin
    
    start = 0
    while start < len(text):
        end = min(start + target_chars, len(text))
        
        if end >= len(text):
            # Last chunk
            chunks.append(text[start:])
            break
        
        # Find a safe split point (space, punctuation, or newline)
        safe_delimiters = [' ', '.', '!', '?', ';', ':', '\n', '\t']
        split_point = end
        
        # Look backwards from target position to find safe delimiter
        for i in range(end - 1, start + target_chars // 2, -1):
            if text[i] in safe_delimiters:
                split_point = i + 1
                break
        
        chunk = text[start:split_point].strip()
        if chunk:
            chunks.append(chunk)
        
        start = split_point
    
    logger.info(f"Character-based splitting created {len(chunks)} chunks")
    return chunks


def _combine_small_chunks(chunks: List[str], max_tokens: int, min_tokens: int) -> List[str]:
    """Combine small chunks together to reach minimum size while respecting maximum."""
    if not chunks:
        return []
    
    combined_chunks = []
    current_combined = ""
    
    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk:
            continue
        
        # Try to combine with current chunk
        potential_combined = current_combined + "\n\n" + chunk if current_combined else chunk
        potential_tokens = estimate_tokens(potential_combined)
        
        if potential_tokens <= max_tokens:
            current_combined = potential_combined
        else:
            # Current combination would be too large
            if current_combined:
                combined_chunks.append(current_combined)
            current_combined = chunk
    
    # Add the last combined chunk
    if current_combined:
        combined_chunks.append(current_combined)
    
    return combined_chunks


class DocumentChunk:
    """Enhanced document chunk with detailed metadata and token counting."""
    
    def __init__(self, text: str, chunk_id: int, source_file: str, 
                 metadata: Dict[str, Any] = None):
        self.text = text.strip()
        self.chunk_id = chunk_id
        self.source_file = source_file
        self.metadata = metadata or {}
        self.word_count = len(self.text.split())
        self.char_count = len(self.text)
        self.token_count = estimate_tokens(self.text)  # Add accurate token counting
        
        # Quality checks
        self.is_valid = self._validate_chunk()
    
    def _validate_chunk(self) -> bool:
        """Validate chunk quality with improved criteria."""
        if not self.text or len(self.text.strip()) < 10:
            return False
        if self.word_count < 5:
            return False
        if self.token_count > 4000:  # Prevent chunks that are too large for most models
            logger.warning(f"Chunk {self.chunk_id} has {self.token_count} tokens, may be too large")
            return False
        # Always return true for content - we'll let the embedding process handle validation
        return True
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert chunk to dictionary with token information."""
        return {
            'text': self.text,
            'chunk_id': self.chunk_id,
            'source_file': self.source_file,
            'word_count': self.word_count,
            'char_count': self.char_count,
            'token_count': self.token_count,  # Include token count in output
            'is_valid': self.is_valid,
            'metadata': self.metadata
        }


class RobustDocumentParser:
    """Robust document parser with advanced chunking and comprehensive error handling."""
    
    def __init__(self, min_chunk_tokens: int = 100, max_chunk_tokens: int = 2000, 
                 target_chunk_tokens: int = 1000):
        """
        Initialize parser with token-based chunking parameters.
        
        Args:
            min_chunk_tokens: Minimum tokens per chunk
            max_chunk_tokens: Maximum tokens per chunk (safety limit)
            target_chunk_tokens: Target tokens per chunk for optimal processing
        """
        self.min_chunk_tokens = min_chunk_tokens
        self.max_chunk_tokens = max_chunk_tokens
        self.target_chunk_tokens = target_chunk_tokens
        
        self.parsing_stats = {
            'total_files_processed': 0,
            'successful_parses': 0,
            'failed_parses': 0,
            'fallback_uses': 0,
            'libraries_used': {},
            'total_tokens_processed': 0,
            'avg_tokens_per_chunk': 0
        }
        
        logger.info(f"Initialized parser with token limits: {min_chunk_tokens}-{max_chunk_tokens} tokens")
        logger.info(f"Target chunk size: {target_chunk_tokens} tokens")
        logger.info(f"Available PDF libraries: {[lib[0] for lib in PDF_LIBRARIES]}")
        
        if TIKTOKEN_AVAILABLE:
            logger.info("Using tiktoken for accurate token counting")
        else:
            logger.info("Using approximate word-based token counting")
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a file with comprehensive error handling and logging.
        
        Args:
            file_path: Path to the file to parse
            
        Returns:
            Dictionary containing extracted text, chunks, and metadata
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file type not supported
            Exception: If all parsing attempts fail
        """
        start_time = datetime.now()
        file_path = Path(file_path)
        
        logger.info(f"Starting parse of file: {file_path}")
        logger.info(f"File size: {file_path.stat().st_size / 1024:.2f} KB")
        
        # File existence and accessibility checks
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.is_file():
            logger.error(f"Path is not a file: {file_path}")
            raise ValueError(f"Path is not a file: {file_path}")
        
        if file_path.stat().st_size == 0:
            logger.error(f"File is empty: {file_path}")
            raise ValueError(f"File is empty: {file_path}")
        
        # Check file permissions
        if not os.access(file_path, os.R_OK):
            logger.error(f"File is not readable: {file_path}")
            raise PermissionError(f"File is not readable: {file_path}")
        
        self.parsing_stats['total_files_processed'] += 1
        
        try:
            # Determine file type and parse accordingly
            suffix = file_path.suffix.lower()
            
            if suffix == '.pdf':
                result = self._parse_pdf_robust(file_path)
            elif suffix in ['.docx', '.doc']:
                result = self._parse_docx_robust(file_path)
            elif suffix == '.txt':
                result = self._parse_txt_file(file_path)
            else:
                logger.error(f"Unsupported file type: {suffix}")
                raise ValueError(f"Unsupported file type: {suffix}")
            
            # Post-processing validation
            result = self._validate_and_enhance_result(result, file_path, start_time)
            
            self.parsing_stats['successful_parses'] += 1
            logger.info(f"Successfully parsed {file_path} in {(datetime.now() - start_time).total_seconds():.2f}s")
            
            return result
            
        except Exception as e:
            self.parsing_stats['failed_parses'] += 1
            logger.error(f"Failed to parse {file_path}: {str(e)}")
            raise
    
    def _parse_pdf_robust(self, file_path: Path) -> Dict[str, Any]:
        """Parse PDF with multiple libraries and fallback mechanisms."""
        if not PDF_LIBRARIES:
            raise ImportError("No PDF parsing libraries available")
        
        logger.info(f"Attempting PDF parse with {len(PDF_LIBRARIES)} available libraries")
        
        last_exception = None
        
        for lib_name, lib_module in PDF_LIBRARIES:
            try:
                logger.info(f"Trying {lib_name} for PDF parsing...")
                
                if lib_name == 'pymupdf':
                    result = self._parse_with_pymupdf(file_path, lib_module)
                elif lib_name == 'pdfplumber':
                    result = self._parse_with_pdfplumber(file_path, lib_module)
                elif lib_name == 'pypdf2':
                    result = self._parse_with_pypdf2(file_path, lib_module)
                else:
                    continue
                
                # Validate result quality
                if self._validate_extraction_result(result):
                    logger.info(f"Successfully parsed PDF with {lib_name}")
                    self.parsing_stats['libraries_used'][lib_name] = self.parsing_stats['libraries_used'].get(lib_name, 0) + 1
                    return result
                else:
                    logger.warning(f"{lib_name} extraction quality too low, trying next library...")
                    continue
                    
            except Exception as e:
                logger.warning(f"Failed to parse PDF with {lib_name}: {str(e)}")
                last_exception = e
                if lib_name != PDF_LIBRARIES[-1][0]:  # Not the last library
                    logger.info(f"Falling back to next PDF library...")
                    self.parsing_stats['fallback_uses'] += 1
                continue
        
        # If we reach here, all libraries failed
        error_msg = f"All PDF parsing libraries failed. Last error: {last_exception}"
        logger.error(error_msg)
        raise Exception(error_msg)
    
    def _parse_with_pymupdf(self, file_path: Path, fitz_module) -> Dict[str, Any]:
        """Parse PDF using PyMuPDF (fitz)."""
        logger.debug("Opening PDF with PyMuPDF...")
        
        text_content = ""
        metadata = {"file_type": "pdf", "pages": 0, "library": "pymupdf"}
        
        try:
            # Open document with error handling
            doc = fitz_module.open(str(file_path))
            metadata["pages"] = doc.page_count
            
            logger.info(f"PDF has {doc.page_count} pages")
            
            for page_num in range(doc.page_count):
                try:
                    page = doc[page_num]
                    page_text = page.get_text()
                    
                    if page_text and len(page_text.strip()) > 0:
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text
                        logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                    else:
                        logger.warning(f"Page {page_num + 1} contains no extractable text")
                        
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            doc.close()
            
        except Exception as e:
            raise Exception(f"PyMuPDF parsing error: {e}")
        
        if not text_content.strip():
            raise Exception("No text content extracted from PDF")
        
        chunks = self._split_into_chunks(text_content, str(file_path))
        
        return {
            "file_path": str(file_path),
            "file_type": "pdf",
            "raw_text": text_content,
            "chunks": [chunk.to_dict() for chunk in chunks],
            "total_chunks": len(chunks),
            "total_words": len(text_content.split()),
            "total_characters": len(text_content),
            "metadata": metadata
        }
    
    def _parse_with_pdfplumber(self, file_path: Path, pdfplumber_module) -> Dict[str, Any]:
        """Parse PDF using pdfplumber."""
        logger.debug("Opening PDF with pdfplumber...")
        
        text_content = ""
        metadata = {"file_type": "pdf", "pages": 0, "library": "pdfplumber"}
        
        try:
            with pdfplumber_module.open(str(file_path)) as pdf:
                metadata["pages"] = len(pdf.pages)
                
                logger.info(f"PDF has {len(pdf.pages)} pages")
                
                for page_num, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        
                        if page_text and len(page_text.strip()) > 0:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text
                            logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                        else:
                            logger.warning(f"Page {page_num + 1} contains no extractable text")
                            
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
                        
        except Exception as e:
            raise Exception(f"pdfplumber parsing error: {e}")
        
        if not text_content.strip():
            raise Exception("No text content extracted from PDF")
        
        chunks = self._split_into_chunks(text_content, str(file_path))
        
        return {
            "file_path": str(file_path),
            "file_type": "pdf", 
            "raw_text": text_content,
            "chunks": [chunk.to_dict() for chunk in chunks],
            "total_chunks": len(chunks),
            "total_words": len(text_content.split()),
            "total_characters": len(text_content),
            "metadata": metadata
        }
    
    def _parse_with_pypdf2(self, file_path: Path, pypdf2_module) -> Dict[str, Any]:
        """Parse PDF using PyPDF2."""
        logger.debug("Opening PDF with PyPDF2...")
        
        text_content = ""
        metadata = {"file_type": "pdf", "pages": 0, "library": "pypdf2"}
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = pypdf2_module.PdfReader(file)
                metadata["pages"] = len(pdf_reader.pages)
                
                logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        
                        if page_text and len(page_text.strip()) > 0:
                            text_content += f"\n--- Page {page_num + 1} ---\n"
                            text_content += page_text
                            logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                        else:
                            logger.warning(f"Page {page_num + 1} contains no extractable text")
                            
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
                        
        except Exception as e:
            raise Exception(f"PyPDF2 parsing error: {e}")
        
        if not text_content.strip():
            raise Exception("No text content extracted from PDF")
        
        chunks = self._split_into_chunks(text_content, str(file_path))
        
        return {
            "file_path": str(file_path),
            "file_type": "pdf",
            "raw_text": text_content,
            "chunks": [chunk.to_dict() for chunk in chunks],
            "total_chunks": len(chunks),
            "total_words": len(text_content.split()),
            "total_characters": len(text_content),
            "metadata": metadata
        }
    
    def _parse_docx_robust(self, file_path: Path) -> Dict[str, Any]:
        """Parse DOCX file with comprehensive error handling."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX parsing")
        
        logger.info("Starting DOCX parsing...")
        
        text_content = ""
        metadata = {"file_type": "docx", "paragraphs": 0, "tables": 0}
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text and paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    metadata["paragraphs"] += 1
                    
            logger.info(f"Extracted {metadata['paragraphs']} paragraphs")
            
            # Extract tables
            for table_num, table in enumerate(doc.tables):
                metadata["tables"] += 1
                table_text = ""
                
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    
                    if row_text:
                        table_text += " | ".join(row_text) + "\n"
                
                if table_text:
                    text_content += f"\n--- Table {table_num + 1} ---\n"
                    text_content += table_text
            
            logger.info(f"Extracted {metadata['tables']} tables")
            
        except Exception as e:
            raise Exception(f"DOCX parsing error: {e}")
        
        if not text_content.strip():
            raise Exception("No text content extracted from DOCX")
        
        chunks = self._split_into_chunks(text_content, str(file_path))
        
        return {
            "file_path": str(file_path),
            "file_type": "docx",
            "raw_text": text_content,
            "chunks": [chunk.to_dict() for chunk in chunks],
            "total_chunks": len(chunks),
            "total_words": len(text_content.split()),
            "total_characters": len(text_content),
            "metadata": metadata
        }
    
    def _parse_txt_file(self, file_path: Path) -> Dict[str, Any]:
        """Parse plain text file."""
        logger.info("Starting TXT file parsing...")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            text_content = ""
            encoding_used = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    encoding_used = encoding
                    logger.info(f"Successfully read TXT file with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    continue
                    
            if not text_content:
                raise Exception("Could not read TXT file with any supported encoding")
            
            # Basic text cleanup
            text_content = text_content.strip()
            
            metadata = {
                "file_type": "txt", 
                "encoding": encoding_used,
                "lines": len(text_content.split('\n'))
            }
            
            logger.info(f"Parsed TXT file: {len(text_content)} characters, {metadata['lines']} lines")
            
            # Split into chunks
            chunks = self._split_into_chunks(text_content, str(file_path))
            
            return {
                "file_path": str(file_path),
                "file_type": "txt",
                "raw_text": text_content,
                "chunks": [chunk.to_dict() for chunk in chunks],
                "total_chunks": len(chunks),
                "total_words": len(text_content.split()),
                "total_characters": len(text_content),
                "metadata": metadata
            }
            
        except Exception as e:
            raise Exception(f"TXT parsing error: {e}")
    
    def _validate_extraction_result(self, result: Dict[str, Any]) -> bool:
        """Validate the quality of text extraction."""
        if not result or not isinstance(result, dict):
            return False
        
        raw_text = result.get('raw_text', '')
        if not raw_text or len(raw_text.strip()) < 50:
            logger.warning("Extracted text too short")
            return False
        
        # Check word count
        word_count = len(raw_text.split())
        if word_count < 10:
            logger.warning(f"Word count too low: {word_count}")
            return False
        
        # Check for mostly gibberish (low alphabetic ratio)
        alpha_ratio = sum(c.isalpha() or c.isspace() for c in raw_text) / len(raw_text)
        if alpha_ratio < 0.5:
            logger.warning(f"Text quality poor, alphabetic ratio: {alpha_ratio:.2f}")
            return False
        
        # Check chunks
        chunks = result.get('chunks', [])
        valid_chunks = [chunk for chunk in chunks if chunk.get('is_valid', False)]
        
        if len(valid_chunks) == 0:
            logger.warning("No valid chunks produced")
            return False
        
        logger.info(f"Extraction validation passed: {len(valid_chunks)}/{len(chunks)} valid chunks")
        return True
    
    def _split_into_chunks(self, text: str, source_file: str) -> List[DocumentChunk]:
        """
        Split text into chunks using advanced recursive splitting with token counting.
        
        This method implements a comprehensive multi-level splitting strategy:
        1. Analyzes text size and applies appropriate splitting strategy
        2. Uses recursive splitting with token counting for accuracy
        3. Ensures all chunks are within token limits
        4. Provides detailed logging of the splitting process
        
        Args:
            text: Text content to split into chunks
            source_file: Path to the source file for metadata
            
        Returns:
            List of DocumentChunk objects with proper token limits
        """
        if not text or not text.strip():
            logger.warning("Empty text provided for chunking")
            return []

        text = text.strip()
        original_tokens = estimate_tokens(text)
        original_words = len(text.split())
        
        logger.info(f"Starting advanced chunking: {len(text)} chars, {original_words} words, ~{original_tokens} tokens")

        # Clean and normalize text first
        text = self._clean_text(text)
        
        # Use recursive splitting with token counting
        try:
            text_chunks = split_text_recursively(
                text=text,
                max_tokens=self.max_chunk_tokens,
                min_tokens=self.min_chunk_tokens
            )
            
            logger.info(f"Recursive splitting produced {len(text_chunks)} raw text chunks")
            
        except Exception as e:
            logger.error(f"Recursive splitting failed: {e}")
            # Fallback to simple word-based splitting
            logger.info("Falling back to simple word-based splitting")
            words = text.split()
            chunk_size = min(self.target_chunk_tokens, 500)  # Conservative fallback
            text_chunks = []
            
            for i in range(0, len(words), chunk_size):
                chunk_words = words[i:i + chunk_size]
                chunk_text = ' '.join(chunk_words)
                text_chunks.append(chunk_text)
        
        # Convert text chunks to DocumentChunk objects
        document_chunks = []
        total_tokens = 0
        
        for i, chunk_text in enumerate(text_chunks):
            chunk = DocumentChunk(
                text=chunk_text,
                chunk_id=i,
                source_file=source_file,
                metadata={
                    'splitting_method': 'recursive_token_based',
                    'original_text_tokens': original_tokens
                }
            )
            
            if chunk.is_valid:
                document_chunks.append(chunk)
                total_tokens += chunk.token_count
                
                # Log warning for chunks that might be problematic
                if chunk.token_count > self.max_chunk_tokens * 0.9:  # 90% of max
                    logger.warning(f"Chunk {i} is large: {chunk.token_count} tokens")
                elif chunk.token_count < self.min_chunk_tokens:
                    logger.debug(f"Chunk {i} is small: {chunk.token_count} tokens")
            else:
                logger.warning(f"Chunk {i} failed validation, skipping")
        
        # Update statistics
        self.parsing_stats['total_tokens_processed'] += total_tokens
        if document_chunks:
            self.parsing_stats['avg_tokens_per_chunk'] = total_tokens / len(document_chunks)
        
        # Log final statistics
        if document_chunks:
            token_counts = [chunk.token_count for chunk in document_chunks]
            logger.info(f"Created {len(document_chunks)} valid chunks")
            logger.info(f"Token distribution - Min: {min(token_counts)}, Max: {max(token_counts)}, "
                       f"Avg: {sum(token_counts)/len(token_counts):.1f}")
            logger.info(f"Total tokens: {total_tokens} (coverage: {(total_tokens/original_tokens)*100:.1f}%)")
            
            # Validation check: ensure no chunk exceeds limits
            oversized_chunks = [c for c in document_chunks if c.token_count > self.max_chunk_tokens]
            if oversized_chunks:
                logger.error(f"Found {len(oversized_chunks)} chunks exceeding token limit!")
                for chunk in oversized_chunks:
                    logger.error(f"  Chunk {chunk.chunk_id}: {chunk.token_count} tokens")
        else:
            logger.error("No valid chunks created from text!")

        return document_chunks
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common PDF artifacts
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
        
        # Fix common encoding issues
        replacements = {
            'â€™': "'",
            'â€œ': '"',
            'â€\x9d': '"',
            'â€"': '-',
            'â€¢': '•',
        }
        
        for old, new in replacements.items():
            text = text.replace(old, new)
        
        return text.strip()
    
    def _validate_and_enhance_result(self, result: Dict[str, Any], file_path: Path, start_time: datetime) -> Dict[str, Any]:
        """Validate and enhance parsing result with additional metadata."""
        
        # Add timing information
        parse_duration = (datetime.now() - start_time).total_seconds()
        result['parse_duration_seconds'] = parse_duration
        result['parsed_at'] = datetime.now().isoformat()
        
        # Add file metadata
        result['file_size_bytes'] = file_path.stat().st_size
        result['file_modified'] = datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
        
        # Validate chunks
        chunks = result.get('chunks', [])
        valid_chunks = [chunk for chunk in chunks if chunk.get('is_valid', True)]
        invalid_chunk_count = len(chunks) - len(valid_chunks)
        
        if invalid_chunk_count > 0:
            logger.warning(f"Discarded {invalid_chunk_count} invalid chunks")
            result['chunks'] = valid_chunks
            result['total_chunks'] = len(valid_chunks)
            result['invalid_chunks_discarded'] = invalid_chunk_count
        
        # Final validation
        if not result.get('raw_text') or not result['raw_text'].strip():
            raise Exception("No valid text content extracted")
        
        if len(valid_chunks) == 0:
            raise Exception("No valid chunks created from document")
        
        # Add quality metrics
        result['extraction_quality'] = {
            'total_characters': len(result['raw_text']),
            'total_words': len(result['raw_text'].split()),
            'valid_chunks': len(valid_chunks),
            'avg_chunk_words': sum(chunk.get('word_count', 0) for chunk in valid_chunks) / len(valid_chunks) if valid_chunks else 0,
            'alphabetic_ratio': sum(c.isalpha() or c.isspace() for c in result['raw_text']) / len(result['raw_text']) if result['raw_text'] else 0
        }
        
        logger.info(f"Extraction quality: {result['extraction_quality']}")
        
        return result
    
    def get_parsing_stats(self) -> Dict[str, Any]:
        """Get parsing statistics."""
        return {
            **self.parsing_stats,
            'success_rate': self.parsing_stats['successful_parses'] / max(1, self.parsing_stats['total_files_processed']),
            'available_libraries': [lib[0] for lib in PDF_LIBRARIES]
        }


# Factory function for backward compatibility
def parse_document(file_path: str, min_chunk_tokens: int = 100, max_chunk_tokens: int = 2000, 
                  target_chunk_tokens: int = 1000, **kwargs) -> Dict[str, Any]:
    """
    Parse a document file with robust error handling and advanced chunking.
    
    Args:
        file_path: Path to the file to parse
        min_chunk_tokens: Minimum tokens per chunk (default: 100)
        max_chunk_tokens: Maximum tokens per chunk (default: 2000, safe for most embedding models)
        target_chunk_tokens: Target tokens per chunk for optimal processing (default: 1000)
        
    Returns:
        Dictionary containing extracted text, chunks, and metadata with token information
    """
    parser = RobustDocumentParser(min_chunk_tokens, max_chunk_tokens, target_chunk_tokens)
    result = parser.parse_file(file_path)
    
    # Add token statistics to result
    if 'chunks' in result:
        token_counts = [chunk.get('token_count', 0) for chunk in result['chunks']]
        result['token_statistics'] = {
            'total_tokens': sum(token_counts),
            'avg_tokens_per_chunk': sum(token_counts) / len(token_counts) if token_counts else 0,
            'min_tokens_per_chunk': min(token_counts) if token_counts else 0,
            'max_tokens_per_chunk': max(token_counts) if token_counts else 0,
            'chunks_over_target': len([t for t in token_counts if t > target_chunk_tokens]),
            'chunks_under_minimum': len([t for t in token_counts if t < min_chunk_tokens])
        }
    
    return result


if __name__ == "__main__":
    # Test the parser
    import sys
    
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
        try:
            result = parse_document(test_file)
            print(f"Successfully parsed: {result['file_path']}")
            print(f"Total chunks: {result['total_chunks']}")
            print(f"Total words: {result['total_words']}")
            print(f"Library used: {result['metadata'].get('library', 'unknown')}")
        except Exception as e:
            print(f"Error: {e}")
    else:
        print("Usage: python robust_document_parser.py <file_path>")
